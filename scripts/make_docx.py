#!/usr/bin/env python3
"""Generate docs/ci_cd_demo_script.docx from docs/ci_cd_demo_script.md."""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt
except ImportError:
    print("Install dependency: python -m pip install python-docx", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "ci_cd_demo_script.md"
DOCX_PATH = ROOT / "docs" / "ci_cd_demo_script.docx"


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def build_docx(md_text: str) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_lines))
                p.style = "Intense Quote"
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
            for ri, row in enumerate(rows):
                for j, cell in enumerate(row):
                    if j < len(headers):
                        table.rows[ri + 1].cells[j].text = cell
            continue

        if stripped.startswith("> "):
            add_rich_paragraph(doc, stripped[2:], style="Intense Quote")
            i += 1
            continue

        if re.match(r"^-\s+\[[ xX]\]", stripped):
            while i < len(lines) and re.match(r"^-\s+\[[ xX]\]", lines[i].strip()):
                item = re.sub(r"^-\s+\[[ xX]\]\s*", "", lines[i].strip())
                doc.add_paragraph(item, style="List Bullet")
                i += 1
            continue

        if stripped.startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                doc.add_paragraph(lines[i].strip()[2:], style="List Bullet")
                i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                text = re.sub(r"^\d+\.\s*", "", lines[i].strip())
                doc.add_paragraph(text, style="List Number")
                i += 1
            continue

        if not stripped:
            i += 1
            continue

        add_rich_paragraph(doc, stripped)
        i += 1

    return doc


def main() -> None:
    if not MD_PATH.exists():
        print(f"Missing source file: {MD_PATH}", file=sys.stderr)
        sys.exit(1)

    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = build_docx(md_text)
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"Created: {DOCX_PATH}")


if __name__ == "__main__":
    main()
